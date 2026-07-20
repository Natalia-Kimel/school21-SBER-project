from __future__ import annotations

import csv
import io
import re
from pathlib import Path


def extract_text(filename: str, data: bytes) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix in {".txt", ".md", ".log"}:
        return data.decode("utf-8", errors="replace")
    if suffix == ".pdf":
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(data))
        parts = []
        for i, page in enumerate(reader.pages, 1):
            text = page.extract_text() or ""
            parts.append(f"Страница {i}\n{text}")
        return "\n\n".join(parts)
    if suffix == ".docx":
        from docx import Document
        doc = Document(io.BytesIO(data))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text.strip() for cell in row.cells))
        return "\n".join(parts)
    if suffix == ".xlsx":
        from openpyxl import load_workbook
        wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        parts: list[str] = []
        for ws in wb.worksheets:
            parts.append(f"Лист: {ws.title}")
            for row in ws.iter_rows(values_only=True):
                values = [str(v).strip() for v in row if v is not None and str(v).strip()]
                if values:
                    parts.append(" | ".join(values))
        return "\n".join(parts)
    if suffix == ".csv":
        text = data.decode("utf-8-sig", errors="replace")
        rows = csv.reader(io.StringIO(text))
        return "\n".join(" | ".join(row) for row in rows)
    raise ValueError("Поддерживаются PDF, DOCX, XLSX, CSV, TXT и MD")


def clean_text(text: str) -> str:
    text = re.sub(r"\x00", "", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()
